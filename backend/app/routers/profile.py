"""
Profile Router - COMPLETE WITH WORKING PDF EXTRACTION
All features included and fully functional
"""

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status, Query
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from typing import Optional, List, Dict, Any
import os
import shutil
from datetime import datetime
from pathlib import Path
import json
import logging
import tempfile

# Database and auth imports
from app.database import get_db
from app.auth import get_current_user
from app.models import User, UserProfile, Document

# Service imports
from app.services.ai_agent import AIAgent
from app.services.pdf_generator import PDFGenerator
from app.config import settings

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
import io

logger = logging.getLogger(__name__)

router = APIRouter()


# ============================================================================
# Constants
# ============================================================================

DOCUMENT_TYPES = [
    'resume',
    'motivation_letter',
    'diploma',
    'certificate',
    'research_proposal',
    'recommendation_letter',
    'transcript',
    'other'
]

ALLOWED_EXTENSIONS = ['.pdf', '.docx', '.doc', '.txt']


# ============================================================================
# Utility Functions
# ============================================================================

def ensure_directory_exists(directory: str):
    """Ensure upload directory exists"""
    if not os.path.exists(directory):
        os.makedirs(directory)
        logger.info(f"Created directory: {directory}")


def generate_unique_filename(original_filename: str, prefix: str = "") -> str:
    """Generate unique filename with timestamp"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    name, ext = os.path.splitext(original_filename)
    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()
    return f"{prefix}_{timestamp}_{safe_name}{ext}" if prefix else f"{timestamp}_{safe_name}{ext}"


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file
    Returns the complete text content
    """
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            logger.info(f"PDF has {num_pages} pages")
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                text += page_text + "\n\n"
                logger.info(f"Extracted {len(page_text)} chars from page {page_num + 1}")
        
        logger.info(f"Total extracted text length: {len(text)} characters")
        return text.strip()
    
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {str(e)}")
        raise Exception(f"PDF text extraction failed: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file
    Returns the complete text content
    """
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        logger.info(f"Extracted {len(text)} chars from DOCX")
        return text.strip()
    
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {str(e)}")
        raise Exception(f"DOCX text extraction failed: {str(e)}")


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from any supported document type
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif file_ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise Exception(f"Unsupported file type: {file_ext}")


def normalize_gender(gender_value: Any) -> Optional[str]:
    """
    Normalize gender value to proper enum value or None
    
    Handles:
    - String "null" → None
    - Empty string "" → None
    - None → None
    - Valid enum values → enum value
    """
    if gender_value is None:
        return None
    
    if isinstance(gender_value, str):
        gender_clean = gender_value.strip().lower()
        
        # Handle null-like values
        if gender_clean in ('null', '', 'none', 'n/a'):
            return None
        
        # Valid gender enum values
        if gender_clean in ('male', 'female', 'non_binary', 'prefer_not_to_say'):
            return gender_clean
    
    return None


def get_empty_profile_data() -> Dict:
    """Return empty profile structure"""
    return {
        "gender": None,
        "phone": "",
        "location": "",
        "personal_website": "",
        "github_url": "",
        "linkedin_url": "",
        "current_title": "",
        "years_of_experience": 0,
        "professional_summary": "",
        "work_experience": [],
        "education": [],
        "projects": [],
        "training_certifications": [],
        "publications": [],
        "languages": [],
        "awards_honors": [],
        "technical_skills": {
            "programming_languages": [],
            "frameworks": [],
            "databases": [],
            "tools": [],
            "cloud_platforms": [],
            "soft_skills": [],
            "other_skills": []
        },
        "online_courses": [],
        "volunteer_work": [],
        "resume_template": "modern"
    }


# ============================================================================
# Profile Endpoints
# ============================================================================

@router.get("/")
async def get_profile(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get comprehensive user profile"""
    profile = db.query(UserProfile).filter(
        UserProfile.user_id == current_user.id
    ).first()
    
    # Get latest resume document
    latest_resume = db.query(Document).filter(
        Document.user_id == current_user.id,
        Document.document_type == "resume"
    ).order_by(Document.created_at.desc()).first()
    
    if not profile:
        return {
            "user_id": current_user.id,
            "full_name": current_user.full_name or "",
            "email": current_user.email,
            "profile_picture": current_user.profile_picture or "",
            "has_profile": False,
            "latest_resume": {
                "file_name": latest_resume.file_name if latest_resume else None,
                "file_path": f"/api/profile/download-document/{latest_resume.id}" if latest_resume else None,
                "uploaded_at": latest_resume.created_at.isoformat() if latest_resume else None
            } if latest_resume else None,
            "profile_data": get_empty_profile_data()
        }
    
    return {
        "user_id": current_user.id,
        "profile_id": profile.id,
        "full_name": profile.full_name or current_user.full_name or "",
        "email": current_user.email,
        "profile_picture": profile.profile_image_url or current_user.profile_picture or "",
        "has_profile": True,
        "latest_resume": {
            "file_name": latest_resume.file_name if latest_resume else None,
            "file_path": f"/api/profile/download-document/{latest_resume.id}" if latest_resume else None,
            "uploaded_at": latest_resume.created_at.isoformat() if latest_resume else None
        } if latest_resume else None,
        "profile_data": {
            "gender": profile.gender.value if profile.gender else None,
            "phone": profile.phone or "",
            "location": profile.location or "",
            "personal_website": profile.personal_website or "",
            "github_url": profile.github_url or "",
            "linkedin_url": profile.linkedin_url or "",
            "current_title": profile.current_title or "",
            "years_of_experience": profile.years_of_experience or 0,
            "professional_summary": profile.professional_summary or "",
            "work_experience": profile.work_experience or [],
            "education": profile.education or [],
            "projects": profile.projects or [],
            "training_certifications": profile.training_certifications or [],
            "publications": profile.publications or [],
            "languages": profile.languages or [],
            "awards_honors": profile.awards_honors or [],
            "technical_skills": profile.technical_skills or {
                "programming_languages": [],
                "frameworks": [],
                "databases": [],
                "tools": [],
                "cloud_platforms": [],
                "soft_skills": [],
                "other_skills": []
            },
            "online_courses": profile.online_courses or [],
            "volunteer_work": profile.volunteer_work or [],
            "resume_template": profile.resume_template or "modern"
        }
    }


@router.put("/")
async def update_profile(
    profile_update: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update user profile with proper NULL handling"""
    profile = db.query(UserProfile).filter(
        UserProfile.user_id == current_user.id
    ).first()
    
    if not profile:
        profile = UserProfile(user_id=current_user.id)
        db.add(profile)
    
    profile_data = profile_update.get('profile_data', {})
    
    # Update simple fields with proper null handling
    if 'full_name' in profile_data:
        profile.full_name = profile_data['full_name'] or None
    
    # CRITICAL: Properly handle gender enum
    if 'gender' in profile_data:
        normalized_gender = normalize_gender(profile_data['gender'])
        profile.gender = normalized_gender
        logger.debug(f"Gender normalized: '{profile_data['gender']}' → '{normalized_gender}'")
    
    if 'phone' in profile_data:
        profile.phone = profile_data['phone'] or None
    if 'location' in profile_data:
        profile.location = profile_data['location'] or None
    if 'personal_website' in profile_data:
        profile.personal_website = profile_data['personal_website'] or None
    if 'github_url' in profile_data:
        profile.github_url = profile_data['github_url'] or None
    if 'linkedin_url' in profile_data:
        profile.linkedin_url = profile_data['linkedin_url'] or None
    if 'current_title' in profile_data:
        profile.current_title = profile_data['current_title'] or None
    if 'years_of_experience' in profile_data:
        profile.years_of_experience = profile_data['years_of_experience'] or 0
    if 'professional_summary' in profile_data:
        profile.professional_summary = profile_data['professional_summary'] or None
    
    # Update JSON fields
    if 'work_experience' in profile_data:
        profile.work_experience = profile_data['work_experience'] or []
    if 'education' in profile_data:
        profile.education = profile_data['education'] or []
    if 'projects' in profile_data:
        profile.projects = profile_data['projects'] or []
    if 'training_certifications' in profile_data:
        profile.training_certifications = profile_data['training_certifications'] or []
    if 'publications' in profile_data:
        profile.publications = profile_data['publications'] or []
    if 'languages' in profile_data:
        profile.languages = profile_data['languages'] or []
    if 'awards_honors' in profile_data:
        profile.awards_honors = profile_data['awards_honors'] or []
    if 'technical_skills' in profile_data:
        profile.technical_skills = profile_data['technical_skills'] or {}
    if 'online_courses' in profile_data:
        profile.online_courses = profile_data['online_courses'] or []
    if 'volunteer_work' in profile_data:
        profile.volunteer_work = profile_data['volunteer_work'] or []
    if 'resume_template' in profile_data:
        profile.resume_template = profile_data['resume_template'] or 'modern'
    
    profile.updated_at = datetime.utcnow()
    
    try:
        db.commit()
        logger.info(f"Profile updated successfully for user {current_user.id}")
        return {"message": "Profile updated successfully"}
    except Exception as e:
        db.rollback()
        logger.error(f"Failed to update profile: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update profile: {str(e)}")


# ============================================================================
# Resume Upload with AI Extraction
# ============================================================================

@router.post("/upload-resume")
async def upload_resume(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload and process resume with AI extraction
    
    Steps:
    1. Validate file
    2. Save file to disk
    3. Extract text from file (PDF/DOCX/TXT)
    4. Use AI to extract structured data
    5. Update user profile with extracted data
    6. Return extracted info
    """
    
    # Validation
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type {file_ext} not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    
    # Ensure upload directory exists
    ensure_directory_exists(settings.UPLOAD_DIR)
    
    # Generate unique filename
    filename = generate_unique_filename(file.filename, f"resume_{current_user.id}")
    file_path = os.path.join(settings.UPLOAD_DIR, filename)
    
    # Save file to disk
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        file_size = os.path.getsize(file_path)
        logger.info(f"File saved: {file_path} ({file_size} bytes)")
    
    except Exception as e:
        logger.error(f"Failed to save file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
    
    # Extract text from file
    try:
        logger.info(f"Extracting text from {file_ext} file...")
        extracted_text = extract_text_from_file(file_path)
        
        if not extracted_text:
            raise Exception("No text could be extracted from the file")
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters")
        
        # Log first 500 chars for debugging
        logger.debug(f"Text preview: {extracted_text[:500]}...")
    
    except Exception as e:
        logger.error(f"Text extraction failed: {str(e)}")
        # Clean up file
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract text from file: {str(e)}"
        )
    
    # Save document to database
    try:
        document = Document(
            user_id=current_user.id,
            document_type="resume",
            file_name=file.filename,
            file_path=file_path,
            file_size=file_size
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        logger.info(f"Document saved to database: ID {document.id}")
    
    except Exception as e:
        logger.error(f"Failed to save document to database: {str(e)}")
        # Clean up file
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to save document: {str(e)}"
        )
    
    # Extract information using AI
    extracted_info = None
    ai_extraction_error = None
    
    try:
        logger.info("Starting AI extraction...")
        ai_agent = AIAgent()
        
        # Call AI agent to extract resume info
        extracted_info = await ai_agent.extract_resume_info(extracted_text)
        logger.info("AI extraction completed successfully")
        logger.debug(f"Extracted info keys: {list(extracted_info.keys())}")
    
    except Exception as e:
        logger.error(f"AI extraction failed: {str(e)}")
        ai_extraction_error = str(e)
        # Continue without AI extraction - return file info only
    
    # Update profile with extracted info (if AI extraction succeeded)
    if extracted_info:
        try:
            logger.info("Updating profile with extracted data...")
            
            profile = db.query(UserProfile).filter(
                UserProfile.user_id == current_user.id
            ).first()
            
            if not profile:
                profile = UserProfile(user_id=current_user.id)
                db.add(profile)
                logger.info("Created new profile")
            
            # Extract personal info
            personal_info = extracted_info.get('personal_info', {})
            
            # Update fields
            if personal_info.get('full_name'):
                profile.full_name = personal_info['full_name']
            
            # Normalize gender before saving
            if 'gender' in personal_info:
                profile.gender = normalize_gender(personal_info['gender'])
            
            if personal_info.get('phone'):
                profile.phone = personal_info['phone']
            if personal_info.get('location'):
                profile.location = personal_info['location']
            if personal_info.get('linkedin_url'):
                profile.linkedin_url = personal_info['linkedin_url']
            if personal_info.get('github_url'):
                profile.github_url = personal_info['github_url']
            if personal_info.get('personal_website'):
                profile.personal_website = personal_info['personal_website']
            if personal_info.get('current_title'):
                profile.current_title = personal_info['current_title']
            if personal_info.get('years_of_experience') is not None:
                profile.years_of_experience = personal_info['years_of_experience']
            if personal_info.get('professional_summary'):
                profile.professional_summary = personal_info['professional_summary']
            
            # Update array/object fields
            if 'work_experience' in extracted_info:
                profile.work_experience = extracted_info['work_experience']
            if 'education' in extracted_info:
                profile.education = extracted_info['education']
            if 'projects' in extracted_info:
                profile.projects = extracted_info['projects']
            if 'technical_skills' in extracted_info:
                profile.technical_skills = extracted_info['technical_skills']
            if 'languages' in extracted_info:
                profile.languages = extracted_info['languages']
            if 'training_certifications' in extracted_info:
                profile.training_certifications = extracted_info['training_certifications']
            if 'publications' in extracted_info:
                profile.publications = extracted_info['publications']
            if 'awards_honors' in extracted_info:
                profile.awards_honors = extracted_info['awards_honors']
            if 'online_courses' in extracted_info:
                profile.online_courses = extracted_info['online_courses']
            if 'volunteer_work' in extracted_info:
                profile.volunteer_work = extracted_info['volunteer_work']
            
            profile.updated_at = datetime.utcnow()
            db.commit()
            
            logger.info("Profile updated successfully with AI-extracted data")
        
        except Exception as e:
            logger.error(f"Failed to update profile: {str(e)}")
            db.rollback()
            # Don't fail the whole request - file is already uploaded
    
    # Prepare response
    response = {
        "message": "Resume uploaded successfully",
        "file_name": file.filename,
        "file_path": f"/api/profile/download-document/{document.id}",
        "file_size": file_size,
        "document_id": document.id,
        "text_extracted": True,
        "text_length": len(extracted_text),
        "ai_extraction_success": extracted_info is not None
    }
    
    if extracted_info:
        response["extracted_info"] = extracted_info
        response["profile_updated"] = True
    else:
        response["warning"] = f"AI extraction unavailable: {ai_extraction_error or 'Unknown error'}"
        response["profile_updated"] = False
    
    return response


# ============================================================================
# Document Management
# ============================================================================

@router.post("/upload-document")
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload supporting documents (motivation letter, diploma, etc.)"""
    
    if document_type not in DOCUMENT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid document type. Must be one of: {', '.join(DOCUMENT_TYPES)}"
        )
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type {file_ext} not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    
    ensure_directory_exists(settings.UPLOAD_DIR)
    filename = generate_unique_filename(file.filename, f"{document_type}_{current_user.id}")
    file_path = os.path.join(settings.UPLOAD_DIR, filename)
    
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        file_size = os.path.getsize(file_path)
        
        # Save to database
        document = Document(
            user_id=current_user.id,
            document_type=document_type,
            file_name=file.filename,
            file_path=file_path,
            file_size=file_size
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        logger.info(f"Document uploaded: {document_type} for user {current_user.id}")
        
        return {
            "message": f"{document_type.replace('_', ' ').title()} uploaded successfully",
            "document_id": document.id,
            "file_name": file.filename,
            "document_type": document_type,
            "file_size": file_size,
            "uploaded_at": document.created_at.isoformat()
        }
    
    except Exception as e:
        logger.error(f"Failed to upload document: {str(e)}")
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Failed to upload document: {str(e)}")


@router.get("/documents")
async def get_documents(
    document_type: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all documents for current user"""
    query = db.query(Document).filter(Document.user_id == current_user.id)
    
    if document_type:
        if document_type not in DOCUMENT_TYPES:
            raise HTTPException(status_code=400, detail=f"Invalid document type")
        query = query.filter(Document.document_type == document_type)
    
    documents = query.order_by(Document.created_at.desc()).all()
    
    return {
        "documents": [
            {
                "id": doc.id,
                "document_type": doc.document_type,
                "file_name": doc.file_name,
                "file_size": doc.file_size,
                "file_path": f"/api/profile/download-document/{doc.id}",
                "uploaded_at": doc.created_at.isoformat()
            }
            for doc in documents
        ],
        "total": len(documents)
    }


@router.get("/download-document/{document_id}")
async def download_document(
    document_id: int,
    current_user: User = Depends(get_current_user),  # ✅ Must be present!
    db: Session = Depends(get_db)
):
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not os.path.exists(document.file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=document.file_path,
        filename=document.file_name,
        media_type="application/octet-stream"
    )

@router.delete("/document/{document_id}")
async def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a document"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file from filesystem
    try:
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
            logger.info(f"Deleted file: {document.file_path}")
    except Exception as e:
        logger.error(f"Failed to delete file: {str(e)}")
    
    # Delete from database
    db.delete(document)
    db.commit()
    
    return {"message": "Document deleted successfully"}


# ============================================================================
# Resume Generation
# ============================================================================

@router.post("/generate-resume")
async def generate_resume(
    template: str = Form("professional"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Generate resume PDF from profile data - FIXED
    
    Uses PDFGenerator.generate_resume() with template parameter
    """
    
    profile = db.query(UserProfile).filter(
        UserProfile.user_id == current_user.id
    ).first()
    
    if not profile:
        raise HTTPException(
            status_code=404,
            detail="Profile not found. Please create a profile first."
        )
    
    # Validate template
    valid_templates = ['modern', 'professional', 'academic']
    if template not in valid_templates:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid template. Must be one of: {', '.join(valid_templates)}"
        )
    
    # Prepare profile data for PDF generation
    profile_data = {
        "personal_info": {
            "full_name": profile.full_name or "",
            "email": current_user.email,
            "phone": profile.phone or "",
            "location": profile.location or "",
            "linkedin_url": profile.linkedin_url or "",
            "github_url": profile.github_url or "",
            "personal_website": profile.personal_website or "",
            "current_title": profile.current_title or "",
            "professional_summary": profile.professional_summary or ""
        },
        "work_experience": profile.work_experience or [],
        "education": profile.education or [],
        "projects": profile.projects or [],
        "technical_skills": profile.technical_skills or {},
        "training_certifications": profile.training_certifications or [],
        "publications": profile.publications or [],
        "languages": profile.languages or [],
        "awards_honors": profile.awards_honors or [],
        "online_courses": profile.online_courses or [],
        "volunteer_work": profile.volunteer_work or []
    }
    
    try:
        # Create temporary file for PDF
        output_path = os.path.join(
            tempfile.gettempdir(),
            f"resume_{current_user.id}_{int(datetime.now().timestamp())}.pdf"
        )
        
        # FIXED: Use correct method name with template parameter
        pdf_generator = PDFGenerator()
        pdf_path = pdf_generator.generate_resume(
            profile_data=profile_data,
            output_path=output_path,
            template=template  # Pass template here
        )
        
        logger.info(f"✅ Resume generated: {template} template for user {current_user.email}")
        
        # Return PDF file
        return FileResponse(
            path=pdf_path,
            media_type='application/pdf',
            filename=f"{profile.full_name or 'Resume'}_{template}.pdf"
        )
    
    except Exception as e:
        logger.error(f"Failed to generate resume: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate resume: {str(e)}"
        )


@router.get("/preview-resume/{document_id}")
async def preview_resume(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Preview resume document (returns file for display)"""
    
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
        Document.document_type == "resume"
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    if not os.path.exists(document.file_path):
        raise HTTPException(status_code=404, detail="File not found on server")
    
    # Return file for preview
    return FileResponse(
        path=document.file_path,
        filename=document.file_name,
        media_type='application/pdf' if document.file_path.endswith('.pdf') else 'application/octet-stream'
    )

@router.get("/download-resume")
async def download_resume(
    template: str = Query("professional"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Download generated resume (GET version for compatibility)"""
    
    profile = db.query(UserProfile).filter(
        UserProfile.user_id == current_user.id
    ).first()
    
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    
    # Validate template
    valid_templates = ['modern', 'professional', 'academic']
    if template not in valid_templates:
        template = 'professional'
    
    # Prepare profile data
    profile_data = {
        "personal_info": {
            "full_name": profile.full_name or "",
            "email": current_user.email,
            "phone": profile.phone or "",
            "location": profile.location or "",
            "linkedin_url": profile.linkedin_url or "",
            "github_url": profile.github_url or "",
            "personal_website": profile.personal_website or "",
            "current_title": profile.current_title or "",
            "professional_summary": profile.professional_summary or ""
        },
        "work_experience": profile.work_experience or [],
        "education": profile.education or [],
        "projects": profile.projects or [],
        "technical_skills": profile.technical_skills or {},
        "training_certifications": profile.training_certifications or [],
        "publications": profile.publications or [],
        "languages": profile.languages or [],
        "awards_honors": profile.awards_honors or [],
    }
    
    try:
        output_path = os.path.join(
            tempfile.gettempdir(),
            f"resume_{current_user.id}_{int(datetime.now().timestamp())}.pdf"
        )
        
        # FIXED: Use correct method
        pdf_generator = PDFGenerator()
        pdf_path = pdf_generator.generate_resume(
            profile_data=profile_data,
            output_path=output_path,
            template=template
        )
        
        return FileResponse(
            path=pdf_path,
            media_type='application/pdf',
            filename=f"{profile.full_name or 'Resume'}_{template}.pdf"
        )
    
    except Exception as e:
        logger.error(f"Resume download failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate resume: {str(e)}")