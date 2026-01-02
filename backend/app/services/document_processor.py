"""
Document Processor Service - FIXED
Handles text extraction from various document formats
"""

import os
from typing import Optional
import logging
from pathlib import Path

# Document processing libraries
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for extracting text from documents"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from document (SYNCHRONOUS - NOT ASYNC)
        
        Args:
            file_path: Path to document file
            
        Returns:
            str: Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        elif file_ext == '.txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is not installed. Install with: pip install PyPDF2")
        
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            
            extracted_text = '\n\n'.join(text)
            logger.info(f"Extracted {len(extracted_text)} characters from PDF")
            return extracted_text
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if Document is None:
            raise ImportError("python-docx is not installed. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            extracted_text = '\n\n'.join(text)
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Extracted {len(text)} characters from TXT")
            return text
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                logger.info(f"Extracted {len(text)} characters from TXT (latin-1)")
                return text
            except Exception as e:
                logger.error(f"TXT extraction failed: {str(e)}")
                raise ValueError(f"Failed to extract text from TXT: {str(e)}")
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")
    
    def validate_document(self, file_path: str, max_size_mb: int = 10) -> bool:
        """
        Validate document file
        
        Args:
            file_path: Path to document
            max_size_mb: Maximum file size in MB
            
        Returns:
            bool: True if valid
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        max_size_bytes = max_size_mb * 1024 * 1024
        
        if file_size > max_size_bytes:
            raise ValueError(f"File too large: {file_size / 1024 / 1024:.2f}MB (max {max_size_mb}MB)")
        
        # Check file extension
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {file_ext}")
        
        return True